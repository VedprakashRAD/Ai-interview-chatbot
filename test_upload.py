import requests
from pathlib import Path
from docx import Document
from reportlab.pdfgen import canvas
import io

def create_test_pdf():
    """Create a test PDF resume"""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)
    c.drawString(100, 750, "SKILLS")
    c.drawString(100, 730, "Python, FastAPI, Machine Learning")
    c.drawString(100, 700, "EXPERIENCE")
    c.drawString(100, 680, "Senior Software Developer")
    c.drawString(100, 650, "EDUCATION")
    c.drawString(100, 630, "Master in Computer Science")
    c.save()
    buffer.seek(0)
    return buffer

def create_test_docx():
    """Create a test DOCX resume"""
    doc = Document()
    doc.add_heading('SKILLS', 0)
    doc.add_paragraph('Python, FastAPI, Machine Learning')
    doc.add_heading('EXPERIENCE', 0)
    doc.add_paragraph('Senior Software Developer')
    doc.add_heading('EDUCATION', 0)
    doc.add_paragraph('Master in Computer Science')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def test_resume_upload():
    url = 'http://127.0.0.1:8001/api/resume/upload'
    
    # Test with PDF
    pdf_buffer = create_test_pdf()
    files = {'file': ('resume.pdf', pdf_buffer, 'application/pdf')}
    response = requests.post(url, files=files)
    print("\nPDF Upload Test:")
    print_response(response)
    
    # Test with DOCX
    docx_buffer = create_test_docx()
    files = {'file': ('resume.docx', docx_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    response = requests.post(url, files=files)
    print("\nDOCX Upload Test:")
    print_response(response)

def print_response(response):
    print(f"Status: {response.status_code}")
    if response.ok:
        data = response.json()
        print("\nParsed Resume:")
        print("\nSkills:")
        for skill in data['skills']:
            print(f"- {skill}")
        print("\nExperience:")
        for exp in data['experience']:
            print(f"- {exp}")
        print("\nEducation:")
        for edu in data['education']:
            print(f"- {edu}")
        print(f"\nDomain: {data['domain']}")
        print(f"Skill Level: {data['skill_level']}")
    else:
        print(f"Error: {response.text}")

if __name__ == "__main__":
    test_resume_upload() 