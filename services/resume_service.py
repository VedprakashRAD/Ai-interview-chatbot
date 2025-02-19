# -*- coding: utf-8 -*-
import re
from typing import BinaryIO, Dict, List
import PyPDF2
from docx import Document
import magic
import io
import os
import fitz  # PyMuPDF
import docx
from uuid import uuid4
from datetime import datetime

class ResumeService:
    ALLOWED_MIME_TYPES = [
        'application/pdf',
        'application/msword',  # .doc
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'  # .docx
    ]
    
    def __init__(self):
        self.nlp = None
        self.web_dev_keywords = {
            'html', 'css', 'javascript', 'react', 'angular', 'vue', 'node', 'frontend', 'backend'
        }
        self.ai_ml_keywords = {
            'machine learning', 'ai', 'deep learning', 'neural networks', 'python', 'tensorflow', 'pytorch'
        }
        try:
            import spacy
            self.nlp = spacy.load("en_core_web_sm")
        except Exception as e:
            print(f"Warning: spaCy not available, using basic parsing: {str(e)}")
        self.upload_dir = "uploads/resumes"
        os.makedirs(self.upload_dir, exist_ok=True)

    def validate_file(self, file_content: bytes) -> bool:
        mime = magic.Magic(mime=True)
        file_type = mime.from_buffer(file_content)
        return file_type in self.ALLOWED_MIME_TYPES

    def save_resume(self, file_content: bytes, original_filename: str) -> str:
        if not self.validate_file(file_content):
            raise ValueError("Invalid file format. Only PDF, DOC, and DOCX files are allowed.")
        
        # Generate unique filename
        ext = os.path.splitext(original_filename)[1]
        filename = f"{uuid4()}{ext}"
        filepath = os.path.join(self.upload_dir, filename)
        
        with open(filepath, "wb") as f:
            f.write(file_content)
        
        return filepath

    def parse_resume(self, filepath: str) -> Dict:
        ext = os.path.splitext(filepath)[1].lower()
        
        if ext == '.pdf':
            return self._parse_pdf(filepath)
        elif ext in ['.doc', '.docx']:
            return self._parse_docx(filepath)
        else:
            raise ValueError("Unsupported file format")

    def _parse_pdf(self, filepath: str) -> Dict:
        text = ""
        with fitz.open(filepath) as doc:
            for page in doc:
                text += page.get_text()
        
        return self._extract_information(text)

    def _parse_docx(self, filepath: str) -> Dict:
        doc = docx.Document(filepath)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return self._extract_information(text)

    def _extract_information(self, text: str) -> Dict:
        # This is a basic implementation - you might want to use more sophisticated
        # NLP techniques or regex patterns for better extraction
        return {
            "skills": self._extract_skills(text),
            "education": self._extract_education(text),
            "experience": self._extract_experience(text),
            "projects": self._extract_projects(text)
        }

    # Add methods to extract specific information
    def _extract_skills(self, text: str) -> List[str]:
        # Implement skill extraction logic
        # This is a placeholder implementation
        return []

    def _extract_education(self, text: str) -> List[Dict]:
        # Implement education extraction logic
        return []

    def _extract_experience(self, text: str) -> List[Dict]:
        # Implement experience extraction logic
        return []

    def _extract_projects(self, text: str) -> List[Dict]:
        # Implement project extraction logic
        return []

    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
                
            return text
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")

    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            return text
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")

    def _split_into_sections(self, text: str) -> Dict[str, List[str]]:
        """Split text into sections"""
        sections = {}
        current_section = None
        current_content = []
        
        # Split text into lines and process each line
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            # Check if this is a section header
            upper_line = line.upper()
            if any(section in upper_line for section in ['SKILLS', 'EXPERIENCE', 'EDUCATION']):
                # Save previous section if exists
                if current_section and current_content:
                    sections[current_section] = current_content
                
                # Start new section
                current_section = next(s for s in ['SKILLS', 'EXPERIENCE', 'EDUCATION'] if s in upper_line)
                current_content = []
            elif current_section:
                current_content.append(line)
        
        # Add the last section
        if current_section and current_content:
            sections[current_section] = current_content
        
        # Process skills section (comma-separated)
        if 'SKILLS' in sections:
            skills_text = ' '.join(sections['SKILLS'])
            sections['SKILLS'] = [s.strip() for s in skills_text.split(',') if s.strip()]
        
        # Limit section sizes
        return {
            'SKILLS': sections.get('SKILLS', [])[:10],
            'EXPERIENCE': sections.get('EXPERIENCE', [])[:5],
            'EDUCATION': sections.get('EDUCATION', [])[:3]
        }

    def _get_section(self, text: str, section_name: str) -> str:
        """Extract section from text"""
        pattern = f"{section_name}.*?(?=\\n\\n|$)"
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            section = match.group(0)
            section = re.sub(f"{section_name}.*?\\n", '', section, flags=re.IGNORECASE)
            return section.strip()
        return ""

    def _basic_extract_skills(self, text):
        """Basic skill extraction without spaCy"""
        common_skills = ['python', 'java', 'javascript', 'html', 'css', 'sql', 'react', 'angular', 'node']
        found_skills = []
        for skill in common_skills:
            if re.search(r'\b' + skill + r'\b', text.lower()):
                found_skills.append(skill)
        return found_skills

    def _basic_extract_experience(self, text):
        """Basic experience extraction without spaCy"""
        experiences = []
        lines = text.split('\n')
        for line in lines:
            if any(word in line.lower() for word in ['work', 'job', 'position', 'experience']):
                experiences.append(line.strip())
        return experiences[:5]

    def _basic_extract_education(self, text):
        """Basic education extraction without spaCy"""
        education = []
        lines = text.split('\n')
        for line in lines:
            if any(word in line.lower() for word in ['degree', 'university', 'college', 'school']):
                education.append(line.strip())
        return education[:3]

    def determine_domain(self, skills: List[str]) -> str:
        """Determine if candidate is web dev or AI/ML"""
        web_dev_count = sum(1 for skill in skills if skill.lower() in self.web_dev_keywords)
        ai_ml_count = sum(1 for skill in skills if skill.lower() in self.ai_ml_keywords)
        
        return 'web_dev' if web_dev_count > ai_ml_count else 'ai_ml'

    def determine_skill_level(self, experience: List[str], skills: List[str]) -> str:
        """Determine candidate's skill level"""
        # Count years of experience
        years = 0
        for exp in experience:
            if 'year' in exp.lower():
                try:
                    years = max(years, int(exp.split()[0]))
                except:
                    continue

        # Basic skill level determination
        if years > 5 or len(skills) > 10:
            return 'expert'
        elif years > 2 or len(skills) > 5:
            return 'intermediate'
        else:
            return 'beginner' 